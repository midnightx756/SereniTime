import docx

doc = docx.Document('databse.docx')

#initialze an empty dictionary
data = {}
for table in doc.tables:
    heading = None
    for para in table.rows[0].cells[0].paragraphs:
        heading = para.text.strip()
        break

    table_data = []

    for row in table.rows[1: ]:
       row_data = [cell.text.strip() for cell in row.cells]
       table_data.append(row_data)

    if heading:
        data[heading]= table_data
    else:
        data[f"Table {len(data)+1}"] = table_data
