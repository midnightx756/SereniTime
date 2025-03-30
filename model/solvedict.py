import docx
import os

def extract_data_from_word_files(file_list):
    """
    Extracts table data from a list of Word files and combines it into a single dictionary.

    Args:
        file_list (list): A list of paths to the Word (.docx) files.

    Returns:
        dict: A dictionary where keys are table headings (or default names)
              and values are lists of lists representing the table data.
    """
    combined_data = {}
    for file_path in file_list:
        try:
            doc = docx.Document(file_path)
            for table in doc.tables:
                heading = None
                for para in table.rows[0].cells[0].paragraphs:
                    heading = para.text.strip()
                    break

                table_data = []
                for row in table.rows[1:]:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                if heading:
                    if heading in combined_data:
                        # If the heading already exists, append the new table data
                        combined_data[heading].extend(table_data)
                    else:
                        combined_data[heading] = table_data
                else:
                    new_key = f"Table from {os.path.basename(file_path)} - {len([k for k in combined_data if k.startswith('Table from ' + os.path.basename(file_path))]) + 1}"
                    combined_data[new_key] = table_data
        except Exception as e:
            print(f"Error processing file '{file_path}': {e}")
    return combined_data

if __name__ == "__main__":
    word_files = [
        'Solutions_Table1.docx',
        'Solutions_Table2.docx',
        'Solutions_Table3.docx',
        'Solutions_Table4.docx',
        'Solutions_table5.docx',
        'Solutions_table6.docx',
        'Solutions_table7.docx',
        'Solutions_table8.docx'
    ]  # Replace with the actual names of your files

    all_data = extract_data_from_word_files(word_files)

    # You can now print or further process the 'all_data' dictionary
    # For example, to see all the extracted data:
    # for heading, data in all_data.items():
    #     print(f"Heading: {heading}")
    #     for row in data:
    #         print(row)
    #     print("-" * 20)

    print("Data from all Word files has been extracted into the 'all_data' dictionary.")